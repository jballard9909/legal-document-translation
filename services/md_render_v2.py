"""
md_render_v2.py

v2 DELTA vs v1 -- CROSS-PAGE TABLE SPLICE
-----------------------------------------
Problem v2 fixes: a table that straddles a source page boundary arrives as two
disconnected fragments. Page N holds the header + first rows; page N+1 holds
the remaining rows. Because pages are translated one call per page, page N+1's
call never saw the header. With the upstream continuation hint in place, page
N+1 now returns HEADERLESS pipe rows ("| label | value |" with no header and
no separator row). v1 could not use them: _looks_like_table() requires a
separator row, so headerless rows fell through to the else-branch and rendered
as justified paragraphs containing literal "|" characters.

v2 adds a splice pass that runs BEFORE any rendering. It reassembles the two
fragments into a single Word table by appending page N+1's leading headerless
rows onto page N's still-open table block.

The splice fires only when ALL of these hold (fail-soft in every other case):
  - page N's last non-footer block IS a table (header + separator present)
  - page N+1's FIRST block is entirely pipe rows with NO separator row
  - the continuation rows' cell count matches the open table's column count
If any check fails, nothing is spliced and behaviour is exactly v1's.

Multi-page spans are handled: if a page's continuation rows consume the whole
page (only footers remain), the table stays open for the next page.

WHAT V2 DOES NOT CHANGE
-----------------------
  - _FOOTER_RE stays Turkish-only, so footer RENDERING is byte-identical to
    v1. The splice uses its own bilingual skip pattern (_SPLICE_SKIP_RE) to
    find trailing footers. NOTE for SETUP_NOTES: because _FOOTER_RE only
    matches "Sayfa N", an English "Page N ..." footer (tr>en direction) still
    renders as a justified body paragraph rather than a centered caption.
    Pre-existing v1 behaviour, deliberately left alone here.
  - _merge_table_continuation_blocks (the stray-blank-line-INSIDE-a-table
    fix) is untouched and still runs per page.
  - No forced page breaks; render_body's flow-naturally decision stands.
  - The [TABLE - STRUCTURE UNCERTAIN, REVIEW] escape hatch is untouched.

Rendering is otherwise identical to v1: render_markdown_block() keeps its
signature and behaviour, now as a thin wrapper over render_blocks().

Stage 5a of document assembly: render the per-page translated Markdown into a
structural-mirror DOCX BODY. Body only -- no affidavit (5b), no source images
(5c), no PDF conversion (5d), no footer/branding. assemble_core.py will import
render_body() and extend the returned Document.

WHY ITS OWN MODULE
------------------
This is where "1:1 mirror" actually happens, so it is factored out and versioned
on its own (same pattern as rasterize.py): independently testable, single
responsibility, the most review-worthy file in the phase. python-docx, to match
the project's Python stack and the FastAPI-wrapper convention.

INPUT
-----
render_body(pages) takes the Aggregate node's page list, each item:
    { "page_index": int,
      "restored_text": "<translated Markdown, PII restored>",
      "missing_placeholders": [...], "unresolved_tokens": [...] }
Pages are sorted by page_index.

GRAMMAR (only what Stage 2 actually emits)
------------------------------------------
    # H1            -> Heading 1                (page titles)
    ## H2           -> Heading 2                (section headings)
    | a | b |       -> real Word table          (needs a :--- separator row)
    [TABLE - STRUCTURE UNCERTAIN, REVIEW]  -> amber review-flag block
    Sayfa N ...     -> centered small caption   (page footer boilerplate)
    (anything else) -> justified body paragraph
Literal clause numbers (1.1., 2.3., I., VIII.) are plain text by design -- Stage
2 was told never to emit Markdown list syntax, so there is no renumbering risk.

FAIL-SOFT (the safety posture, consistent with the rest of the pipeline)
------------------------------------------------------------------------
A |-delimited block with a separator row renders as a Word table. If its rows
are ragged (cell counts disagree with the header) it STILL renders, but an amber
review flag is placed above it -- never a crash, never silent. A |-block with no
separator row is not treated as a table at all; it falls through to paragraphs.
Known OCR/columnar artifacts (fused caption, garbled cells) are reproduced
faithfully and NOT "fixed": guessing is the confident-wrong failure mode.

100% synthetic in the self-test. No real PII.
"""

import re

from docx import Document
from docx.shared import Inches, Pt, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# --- fixed literals / style constants ---
REVIEW_MARKER = "[TABLE \u2014 STRUCTURE UNCERTAIN, REVIEW]"  # em-dash inside
AMBER = "FFF2CC"          # review-flag fill
HEADER_GRAY = "D9D9D9"    # table header fill
USABLE_WIDTH_TWIPS = 9360  # 6.5" at 1440 twips/inch (US Letter minus 1" margins)
MIN_COL_TWIPS = 720        # 0.5" column floor

# A Markdown table separator row: pipes, dashes, optional colons, spaces.
_SEP_RE = re.compile(r"^\s*\|?[\s:|-]*-[\s:|-]*\|?\s*$")
_FOOTER_RE = re.compile(r"^Sayfa\s+\d+")
# Used ONLY by the cross-page splice to skip trailing page-footer blocks when
# looking for a page's still-open table. Bilingual on purpose: the splice must
# work in both directions. Deliberately NOT merged into _FOOTER_RE, which
# governs footer *rendering* -- widening that would change v1 output.
_SPLICE_SKIP_RE = re.compile(
    r"^\s*(Sayfa|Page)\s+\d+|^\s*(FABRICATED DOCUMENT|YAZILIM TEST)",
    re.IGNORECASE)


# ---------------------------------------------------------------------------
# low-level helpers
# ---------------------------------------------------------------------------
def _apply_shading(pr_element, fill):
    """Append a <w:shd> (clear pattern) to a pPr or tcPr element."""
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pr_element.append(shd)


def _shade_paragraph(paragraph, fill):
    _apply_shading(paragraph._p.get_or_add_pPr(), fill)


def _shade_cell(cell, fill):
    _apply_shading(cell._tc.get_or_add_tcPr(), fill)


def _set_base_styles(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(11)


def _set_page(doc):
    for section in doc.sections:
        section.page_width = Inches(8.5)     # US Letter, not the A4 default
        section.page_height = Inches(11)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)


def _add_body_paragraph(doc, text, align=WD_ALIGN_PARAGRAPH.JUSTIFY):
    p = doc.add_paragraph(text)
    p.alignment = align
    return p


def _add_heading(doc, text, level):
    # Built-in heading styles (keeps outline/TOC working), forced black + bold
    # so it reads as a legal heading rather than the themed blue default.
    # Centered: both the document/page titles and the Roman-numeral section
    # headings are centered in decree convention.
    p = doc.add_heading("", level=level)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = RGBColor(0, 0, 0)
    return p


def _add_review_flag(doc, marker_text):
    p = doc.add_paragraph()
    _shade_paragraph(p, AMBER)
    run = p.add_run(marker_text)
    run.bold = True
    return p


def _render_footer(doc, lines):
    for line in lines:
        if line.strip():
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(line.strip())
            run.italic = True
            run.font.size = Pt(8)


# ---------------------------------------------------------------------------
# tables
# ---------------------------------------------------------------------------
def _looks_like_table(lines):
    return (len(lines) >= 2
            and lines[0].lstrip().startswith("|")
            and _SEP_RE.match(lines[1]) is not None)


def _split_row(line):
    parts = line.split("|")
    if parts and parts[0].strip() == "":
        parts = parts[1:]
    if parts and parts[-1].strip() == "":
        parts = parts[:-1]
    return [c.strip() for c in parts]


def _column_widths(header, body, ncols):
    maxlen = [len(header[i]) if i < len(header) else 1 for i in range(ncols)]
    for row in body:
        for i in range(ncols):
            if i < len(row):
                maxlen[i] = max(maxlen[i], len(row[i]))
    total = sum(maxlen) or 1
    return [max(int(USABLE_WIDTH_TWIPS * m / total), MIN_COL_TWIPS)
            for m in maxlen]


def _render_table(doc, lines):
    """Render a Markdown table block. Returns True if ragged (review-flagged)."""
    header = _split_row(lines[0])
    body = [_split_row(l) for l in lines[2:] if l.strip()]
    ncols = len(header)
    ragged = any(len(r) != ncols for r in body)

    if ragged:
        _add_review_flag(
            doc,
            "[TABLE ROWS UNEVEN \u2014 auto-aligned, REVIEW cell placement]")

    table = doc.add_table(rows=1, cols=ncols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    hdr_cells = table.rows[0].cells
    for i in range(ncols):
        hdr_cells[i].text = header[i] if i < len(header) else ""
        _shade_cell(hdr_cells[i], HEADER_GRAY)
        for para in hdr_cells[i].paragraphs:
            for run in para.runs:
                run.bold = True

    for row in body:
        cells = table.add_row().cells
        for i in range(ncols):
            cells[i].text = row[i] if i < len(row) else ""

    widths = _column_widths(header, body, ncols)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            if i < len(widths):
                cell.width = Twips(widths[i])
    return ragged


# ---------------------------------------------------------------------------
# block dispatch
# ---------------------------------------------------------------------------
def _is_pipe_row(line):
    s = line.strip()
    return s.startswith("|") and s.endswith("|") and len(s) > 1


def _merge_table_continuation_blocks(blocks):
    """Guard against a stray blank line INSIDE a Markdown table (between the
    separator row and a data row, or between two data rows -- a known LLM
    Markdown quirk) fragmenting one table into multiple blocks. If a block
    ENDS with a table-row-shaped line and the next block STARTS with one,
    they are almost certainly the same table split by a spurious blank line
    -- merge them back into a single block before block-dispatch runs, so the
    whole table renders as one Word table instead of "first row as a table,
    remaining rows as plain paragraphs" (the exact symptom observed in
    production).
    """
    merged = []
    for block in blocks:
        if merged:
            prev_lines = [l for l in merged[-1].split("\n") if l.strip()]
            cur_lines = [l for l in block.split("\n") if l.strip()]
            if (prev_lines and cur_lines
                    and _is_pipe_row(prev_lines[-1])
                    and _is_pipe_row(cur_lines[0])):
                merged[-1] = merged[-1].rstrip("\n") + "\n" + block.lstrip("\n")
                continue
        merged.append(block)
    return merged


# ---------------------------------------------------------------------------
# v2: cross-page table splice
# ---------------------------------------------------------------------------
def _split_blocks(md):
    """One page's Markdown -> block list, with v1's intra-page merge applied."""
    return _merge_table_continuation_blocks(re.split(r"\n\s*\n", md.strip()))


def _block_lines(block):
    return [l for l in block.split("\n") if l.strip()]


def _is_footerish(block):
    """Trailing page-boilerplate the splice steps over. Splice-only -- has no
    effect on how footers render (see _FOOTER_RE note in the module docstring).
    """
    lines = _block_lines(block)
    return bool(lines) and all(_SPLICE_SKIP_RE.match(l) for l in lines)


def _is_headerless_continuation(block):
    """True for a block that is nothing but pipe rows with NO separator row --
    the shape the upstream continuation hint asks page N+1 to emit.

    A block WITH a separator row is a self-contained table (v1 handles it) and
    must not be spliced: splicing it would silently drop its header.
    """
    lines = _block_lines(block)
    if not lines:
        return False
    if any(_SEP_RE.match(l) for l in lines):
        return False
    return all(_is_pipe_row(l) for l in lines)


def _trailing_open_table_index(blocks):
    """Index of the page's last non-footer block if that block is a table,
    else None. A page whose last content is prose has no open table.
    """
    for i in range(len(blocks) - 1, -1, -1):
        if _is_footerish(blocks[i]):
            continue
        return i if _looks_like_table(_block_lines(blocks[i])) else None
    return None


def _table_ncols(block):
    return len(_split_row(_block_lines(block)[0]))


def _splice_cross_page_tables(pages_blocks):
    """Reassemble tables split across source pages, in place.

    pages_blocks: list (page order) of per-page block lists.

    Carries a reference to the still-open trailing table forward across pages.
    Column count is the guard: continuation rows whose cell count disagrees
    with the open table are left alone, so a genuinely different table that
    happens to start a page is never absorbed.

    Returns (pages_blocks, spliced_row_count) -- the count is for the
    self-test and for callers that want to log that a splice happened.
    """
    open_blocks, open_idx, spliced = None, None, 0

    for blocks in pages_blocks:
        if open_blocks is not None and blocks:
            ncols = _table_ncols(open_blocks[open_idx])
            taken = []
            while blocks and _is_headerless_continuation(blocks[0]):
                rows = _block_lines(blocks[0])
                if any(len(_split_row(r)) != ncols for r in rows):
                    break  # different shape -> not this table's continuation
                taken.extend(rows)
                blocks.pop(0)
            if taken:
                open_blocks[open_idx] = (
                    open_blocks[open_idx].rstrip("\n") + "\n" + "\n".join(taken))
                spliced += len(taken)

        idx = _trailing_open_table_index(blocks)
        if idx is not None:
            open_blocks, open_idx = blocks, idx
        elif not all(_is_footerish(b) for b in blocks):
            # real content that isn't a table closed the table
            open_blocks, open_idx = None, None
        # page left with only footers -> table stays open for the next page

    return pages_blocks, spliced


def render_markdown_block(md, doc, is_first_page=False, page_break_before=False):
    """Parse one page's Markdown and append its elements to doc.

    page_break_before: if True, the break is attached to the FIRST element
    this call renders (heading, paragraph, or table) rather than inserted as
    a separate floating paragraph. Retained for callers that want a forced
    break (e.g. testing); render_body itself no longer requests one between
    source pages -- see render_body's docstring for why.

    v2: kept as a thin wrapper so any caller passing raw Markdown behaves
    exactly as in v1. render_body now pre-splits pages into blocks (so the
    cross-page splice can run across them) and calls render_blocks directly.
    """
    render_blocks(_split_blocks(md), doc,
                  is_first_page=is_first_page,
                  page_break_before=page_break_before)


def render_blocks(blocks, doc, is_first_page=False, page_break_before=False):
    """Render an already-split (and possibly spliced) block list into doc.

    Body identical to v1's render_markdown_block after its split step.
    """
    break_pending = page_break_before

    def _consume_break():
        """Return True exactly once per page: the caller applies the break
        to whatever it just created, then this always returns False after."""
        nonlocal break_pending
        if break_pending:
            break_pending = False
            return True
        return False

    # Caption zone: on the title page, the body blocks that precede the first
    # section/title heading are the case caption (court / parties / "v."). Center
    # them for a faithful decree caption. SCOPED to the first page so ordinary
    # body clauses on later pages -- which can also sit before a heading, e.g.
    # "1.7." ahead of "## II." -- are never mistaken for caption text.
    first_h2_idx = None
    for i, b in enumerate(blocks):
        if b.strip().split("\n")[0].strip().startswith("## "):
            first_h2_idx = i
            break

    for i, block in enumerate(blocks):
        block = block.strip("\n")
        if not block.strip():
            continue
        lines = block.split("\n")
        first = lines[0].strip()
        in_caption = (is_first_page and first_h2_idx is not None
                      and i < first_h2_idx)

        if first.startswith(REVIEW_MARKER):
            p = _add_review_flag(doc, first)
            if _consume_break():
                p.paragraph_format.page_break_before = True
            for line in lines[1:]:
                if line.strip():
                    _add_body_paragraph(doc, line.strip())
        elif _looks_like_table(lines):
            table = _render_table(doc, lines)
            if _consume_break():
                # Tables have no page_break_before of their own -- the break
                # goes on the paragraph inside the table's first cell, which
                # is always present (python-docx guarantees at least one
                # paragraph per cell).
                table_obj = doc.tables[-1]
                table_obj.rows[0].cells[0].paragraphs[0]\
                    .paragraph_format.page_break_before = True
        elif len(lines) == 1 and first.startswith("## "):
            p = _add_heading(doc, first[3:].strip(), 2)
            if _consume_break():
                p.paragraph_format.page_break_before = True
        elif len(lines) == 1 and first.startswith("# "):
            p = _add_heading(doc, first[2:].strip(), 1)
            if _consume_break():
                p.paragraph_format.page_break_before = True
        elif _FOOTER_RE.match(first):
            _render_footer(doc, lines)
        else:
            align = (WD_ALIGN_PARAGRAPH.CENTER if in_caption
                     else WD_ALIGN_PARAGRAPH.JUSTIFY)
            first_para_of_block = True
            for line in lines:
                if line.strip():
                    p = _add_body_paragraph(doc, line.strip(), align=align)
                    if first_para_of_block and _consume_break():
                        p.paragraph_format.page_break_before = True
                    first_para_of_block = False


def render_body(pages):
    """Render all pages into a single mirror-body Document, in page order.

    DESIGN REVERSAL (from the original Stage 5a decision): earlier builds
    forced a hard page break between every source page, for page-to-page
    cross-reference (translation page N <-> source page N). In practice this
    made translation quality on ANY one page a load-bearing input to layout
    on EVERY later page: EN<->TR length mismatch (or a rendering hiccup, e.g.
    a mis-parsed table taking more vertical room than it should) pushes
    content past its natural page boundary, collides with the next forced
    break, and stranded a section alone on a near-empty page -- confirmed in
    production. Content now flows naturally; translation and source page
    numbers may drift out of alignment as a result, which is the accepted
    tradeoff. The "Sayfa N ..." caption already present in each page's own
    text remains the visible marker of where one source page's content ends
    and the next begins -- just inline, not as a hard break.
    """
    doc = Document()
    _set_base_styles(doc)
    _set_page(doc)
    ordered = sorted(pages, key=lambda p: p.get("page_index", 0))

    # v2: split every page first, splice tables across page boundaries, then
    # render. Splitting up front is what makes the cross-page view possible --
    # v1 rendered each page in isolation and so could never see the other half
    # of a split table.
    pages_blocks = [_split_blocks(p.get("restored_text", "")) for p in ordered]
    pages_blocks, _ = _splice_cross_page_tables(pages_blocks)

    for idx, blocks in enumerate(pages_blocks):
        render_blocks(blocks, doc, is_first_page=(idx == 0))
    return doc


# ---------------------------------------------------------------------------
# SELF-TEST -- synthetic only, behind __main__.
# ---------------------------------------------------------------------------
if __name__ == "__main__":
    import os

    # -- Part 1: splice logic in isolation (no docx, fast, precise) ----------
    def _blocks(*mds):
        return [_split_blocks(m) for m in mds]

    OPEN_TABLE = ("| Item | Amount / Detail |\n| :--- | :--- |\n"
                  "| Monthly base child support obligation | $1,240.00 |")
    FOOTER_TR = "Sayfa 2 \u2014 Synthetic Test Data Only"
    CONT_2COL = ("| Payor | [PERSON_1] |\n| Payee | [PERSON_2] |\n"
                 "| Effective date | July 1, 2025 |")

    # fires: open table + matching 2-col headerless continuation
    pb, n = _splice_cross_page_tables(
        _blocks(f"Body text.\n\n{OPEN_TABLE}\n\n{FOOTER_TR}",
                f"{CONT_2COL}\n\n## V. NEXT SECTION"))
    assert n == 3, f"expected 3 rows spliced, got {n}"
    assert len(_block_lines(pb[0][1])) == 6, "rows not appended to open table"
    assert not _is_headerless_continuation(pb[1][0]), "rows not removed"

    # does NOT fire: column count mismatch (a different table)
    _, n = _splice_cross_page_tables(
        _blocks(f"{OPEN_TABLE}\n\n{FOOTER_TR}",
                "| a | b | c |\n| d | e | f |"))
    assert n == 0, "spliced a table with a different column count"

    # does NOT fire: page N ends in prose, no open table
    _, n = _splice_cross_page_tables(
        _blocks(f"{OPEN_TABLE}\n\nClosing prose.\n\n{FOOTER_TR}", CONT_2COL))
    assert n == 0, "spliced onto a table already closed by prose"

    # does NOT fire: next page starts with a SELF-CONTAINED table (has a
    # separator row) -- splicing would silently drop its header
    _, n = _splice_cross_page_tables(
        _blocks(f"{OPEN_TABLE}\n\n{FOOTER_TR}",
                "| H1 | H2 |\n| :--- | :--- |\n| x | y |"))
    assert n == 0, "absorbed a self-contained table and dropped its header"

    # fires across THREE pages: middle page is continuation rows + footer only
    _, n = _splice_cross_page_tables(
        _blocks(f"{OPEN_TABLE}\n\n{FOOTER_TR}",
                f"{CONT_2COL}\n\n{FOOTER_TR}",
                "| Payment method | Income withholding order |\n\n## V. NEXT"))
    assert n == 4, f"3-page span: expected 4 rows spliced, got {n}"

    # bilingual: English footer must also be skipped when finding the table
    _, n = _splice_cross_page_tables(
        _blocks(f"{OPEN_TABLE}\n\nPage 2 \u2014 Synthetic Test Data Only",
                CONT_2COL))
    assert n == 3, "English footer blocked the splice"

    print("[splice] all logic assertions PASS")

    # -- Part 2: end-to-end render ------------------------------------------
    synthetic_pages = [
        {
            "page_index": 0,
            "restored_text": (
                "# SYNTHETIC COURT OF TESTING\n\n"
                "Party A, Case No.: TEST-0001 Petitioner,\n\n"
                "## I. FINDINGS\n\n"
                "1.1. This is a justified body paragraph that should wrap and "
                "read as a single flowing block of legal prose.\n\n"
                "| Item | Amount / Detail |\n| :--- | :--- |\n"
                "| Monthly base child support obligation | $1,240.00 |\n\n"
                "Sayfa 1 \u2014 Synthetic Test Data Only"
            ),
        },
        {
            # the split-table case: headerless continuation rows lead the page
            "page_index": 1,
            "restored_text": (
                "| Payor | [PERSON_1] |\n"
                "| Payee | [PERSON_2] |\n"
                "| Effective date | July 1, 2025 |\n"
                "| Payment method | Income withholding order |\n"
                "| Health insurance coverage | Maintained by Respondent |\n"
                "| Uninsured medical/dental expenses | Shared 5596 / 4596 |\n\n"
                "## II. SIGNATURES\n\n"
                "[TABLE \u2014 STRUCTURE UNCERTAIN, REVIEW]\n"
                "Attorney One, Attorney Two\n\n"
                "A ragged table follows:\n\n"
                "| Col A | Col B | Col C |\n| :--- | :--- | :--- |\n"
                "| only two | cells |\n| three | proper | cells |\n\n"
                "Sayfa 2 \u2014 Synthetic Test Data Only"
            ),
        },
    ]

    doc = render_body(synthetic_pages)
    tables = doc.tables
    print(f"[render] paragraphs={len(doc.paragraphs)} tables={len(tables)}")

    # THE regression this file exists for: the split table is ONE table,
    # header + 7 data rows, not a 2-row table plus six stray paragraphs.
    assert len(tables) == 2, f"expected 2 tables (spliced + ragged), got {len(tables)}"
    spliced = tables[0]
    assert len(spliced.rows) == 8, \
        f"split table should be 1 header + 7 data rows, got {len(spliced.rows)}"
    assert len(spliced.columns) == 2, "split table should be 2 columns"
    assert spliced.rows[1].cells[0].text.startswith("Monthly base"), \
        "first data row lost"
    assert spliced.rows[7].cells[0].text.startswith("Uninsured"), \
        "last continuation row lost"

    # no literal pipes leaked into body paragraphs
    assert not any(p.text.strip().startswith("|") for p in doc.paragraphs), \
        "a pipe row rendered as a paragraph -- splice missed rows"

    # v1 behaviour preserved: review-marker block + ragged-table flag
    shaded = sum(
        1 for p in doc.paragraphs
        if p._p.find(qn("w:pPr")) is not None
        and p._p.pPr.find(qn("w:shd")) is not None
    )
    print(f"[render] amber-shaded paragraphs={shaded}")
    assert shaded == 2, f"expected 2 review flags, got {shaded}"

    out = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                       "md_render_v2_selftest.docx")
    doc.save(out)
    print(f"[render] self-test doc saved -> {out}")
    print("[render] SELF-TEST PASS")